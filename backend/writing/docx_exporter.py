"""
DOCX export for writing module using python-docx.
Converts markdown content to a formatted Word document.
"""
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def markdown_to_docx(content: str, title: str, output_path: str) -> str:
    """
    Convert markdown content to a DOCX file.
    Returns the output file path.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise ImportError("python-docx is required: pip install python-docx")

    doc = Document()

    # Document styling
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Title
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.color.rgb = RGBColor(0x02, 0x8a, 0x39)  # #028a39 green

    doc.add_paragraph()  # spacer

    lines = content.split('\n')

    for line in lines:
        line = line.rstrip()

        if not line:
            doc.add_paragraph()
            continue

        # Heading levels
        if line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
            p.runs[0].font.color.rgb = RGBColor(0x02, 0x8a, 0x39)
        elif line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
            p.runs[0].font.color.rgb = RGBColor(0x02, 0x8a, 0x39)
        elif line.startswith('- ') or line.startswith('* '):
            # Bullet point
            p = doc.add_paragraph(style='List Bullet')
            _add_inline_formatted_text(p, line[2:])
        elif re.match(r'^\d+\. ', line):
            # Numbered list
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\. ', '', line)
            _add_inline_formatted_text(p, text)
        else:
            p = doc.add_paragraph()
            _add_inline_formatted_text(p, line)

    # Save document
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    logger.info(f"DOCX saved to {output_path}")
    return output_path


def _add_inline_formatted_text(paragraph, text: str):
    """Add text to paragraph, handling **bold** and *italic* markdown."""
    import re
    # Pattern: **bold**, *italic*, `code`
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    parts = pattern.split(text)

    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
        else:
            paragraph.add_run(part)
